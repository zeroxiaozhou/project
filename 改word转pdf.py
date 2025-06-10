import os
import win32com.client
import tqdm
import pandas
import docx
from tqdm import trange
import shutil
from datetime import datetime
import openpyxl
from openpyxl.styles import Alignment, Font
import pandas as pd
import random
# 获取当前日期
today = datetime.now()
# 格式化日期为MM-DD
formatted_date = today.strftime('%m-%d')


def input_word():
    df = pandas.read_excel(rf'E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}\喜行广州注册报备{formatted_date}.xlsx')
    file = docx.Document(r'C:\Users\XYCX\Desktop\喜行约车协议.docx')
    for i in trange(len(df)):
        name = df['姓名'][i]
        id = df['身份证号'][i]
        ID = df['司机ID'][i]
        file.paragraphs[2].add_run(name)
        file.paragraphs[3].add_run(id)
        file.save(rf'E:\喜行约车司机协议\25年喜行约车协议{formatted_date}\{ID}.docx')
        file = docx.Document(r'C:\Users\XYCX\Desktop\喜行约车协议.docx')


def word_switch_pdf():
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False  # 不显示Word窗口

    source_dir = rf'E:\喜行约车司机协议\25年喜行约车协议{formatted_date}'
    target_dir = rf'E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}\PDF总数'

    for name in tqdm.tqdm(os.listdir(source_dir)):
        docx_path = os.path.join(source_dir, name)
        pdf_path = os.path.join(target_dir, name.split('.')[0] + '.pdf')
        # if name.lower().endswith('.docx'):
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
    word.Quit()


def copy_file():
    # 文件路径
    excel_file_path = fr"E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}\喜行广州注册报备{formatted_date}.xlsx"
    source_folder = fr"E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}\PDF总数"

    # 读取 Excel 表格
    df = pd.read_excel(excel_file_path)

    # 清理司机ID数据
    df['司机ID'] = df['司机ID'].astype(str).str.strip()

    # 指定要创建的文件夹名称列表
    folders = ['clx', 'hxt', 'lh', 'llh', 'll', 'liangli', 'lxr', 'lly', 'xxk', 'zjq']

    # 创建每个文件夹
    for folder in folders:
        os.makedirs(os.path.join(rf"E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}", folder),
                    exist_ok=True)
        # exist_ok=True 确保如果文件夹已存在不会报错
    print("文件夹创建完成！")

    # 获取源文件夹中的 PDF 文件
    pdf_files = [f for f in os.listdir(source_folder) if f.endswith('.pdf')]

    # 根据司机ID平均分配PDF文件
    file_assignments = {folder: [] for folder in folders}

    for i, pdf_file in enumerate(pdf_files):
        target_folder = folders[i % len(folders)]  # 循环选择文件夹
        file_assignments[target_folder].append(pdf_file)

    # 为每个文件夹生成 Excel 表格
    for folder in folders:
        folder_path = os.path.join(rf"E:\喜行约车司机协议PDF分类\25年喜行约车协议PDF{formatted_date}", folder)
        assigned_files = file_assignments[folder]

        # 提取与当前文件夹分配的 PDF 文件对应的司机 ID
        assigned_driver_ids = [pdf_file.split('.')[0].strip() for pdf_file in assigned_files]

        # 从原 DataFrame 中筛选出对应的行
        assigned_df = df[df['司机ID'].isin(assigned_driver_ids)].copy()

        # 生成新的 Excel 文件
        new_excel_file_path = os.path.join(folder_path, f"{folder}.xlsx")
        assigned_df.to_excel(new_excel_file_path, index=False)

        # 复制 PDF 文件到相应文件夹
        for pdf_file in assigned_files:
            shutil.copy(os.path.join(source_folder, pdf_file), folder_path)


def input_word2():
    # 读取Excel文件
    df = pd.read_excel(rf"C:\Users\XYCX\Desktop\名单.xlsx")

    # 随机选择Word模板文件
    file = docx.Document(rf"C:\Users\XYCX\Desktop\1.3.4驾驶员绩效考核表{random.randint(1, 4)}.docx")

    for i in trange(len(df)):
        name = df['姓名'][i]
        driver = df['内岗证号'][i]

        # 获取第18段落
        paragraph = file.paragraphs[18]

        # 查找"姓名："和"内岗证号："的位置
        text = paragraph.text
        name_pos = text.find("姓名：")
        name_pos2 = text.find("内岗证号：")

        # 如果找到"姓名："和"内岗证号："，就插入相应内容
        if name_pos != -1 and name_pos2 != -1:
            # 获取"姓名："后的内容直到"内岗证号："
            before_name = text[:name_pos + len("姓名：")]
            after_name = text[name_pos + len("姓名："):name_pos2]

            # 获取"内岗证号："后的内容
            before_name2 = text[:name_pos2 + len("内岗证号：")]
            after_name2 = text[name_pos2 + len("内岗证号："):]

            # 只插入姓名和内岗证号的部分，避免重复
            new_text = before_name + name + after_name  # 在"姓名："后插入姓名
            new_text2 = before_name2 + str(driver) + after_name2  # 在"内岗证号："后插入内岗证号

            # 清除原段落内容
            paragraph.clear()

            # 添加新的姓名和内岗证号内容
            paragraph.add_run(new_text)  # 添加更新后的姓名部分
            paragraph.add_run(new_text2)  # 添加更新后的内岗证号部分

            # 保存修改后的Word文档
            file.save(rf"C:\Users\XYCX\Desktop\绩效考核\{name}.docx")

        # 每次循环后重新加载随机的Word模板
        file = docx.Document(rf"C:\Users\XYCX\Desktop\1.3.4驾驶员绩效考核表{random.randint(1, 4)}.docx")


def input_word3():
    # 读取Word文档
    for i in os.listdir(rf"C:\Users\XYCX\Desktop\绩效考核"):
        df = docx.Document(rf"C:\Users\XYCX\Desktop\绩效考核\{i}")

        # 获取指定段落的文本
        x = df.paragraphs[18].text

        # 找到第一个“姓名：”的索引
        first_name_index = x.find("姓名：")

        # 找到第二个“姓名：”的索引
        second_name_index = x.find("姓名：", first_name_index + 1)

        # 如果找到了第二个“姓名：”，则将其删除
        if second_name_index != -1:
            x = x[:second_name_index] + x[second_name_index + len("姓名："):]
        # 更新该段落的文本
        df.paragraphs[18].clear()  # 清空原来的段落
        df.paragraphs[18].add_run(x)  # 添加修改后的文本
        df.save(rf"C:\Users\XYCX\Desktop\新绩效考核\{i}")


def input_word4():
    """服务质量信誉档案"""
    # 读取第一张表
    df1 = pd.read_excel(r"C:\Users\XYCX\Desktop\海口市数据1.xlsx")
    # 打开Excel文件
    workbook = openpyxl.load_workbook(r"C:\Users\XYCX\Desktop\附件4：网约出租汽车驾驶员服务质量信誉档案.xlsx")
    # 获取第二个工作表
    second_sheet = workbook.worksheets[0]
    second_sheet1 = workbook.worksheets[1]
    second_sheet2 = workbook.worksheets[2]
    second_sheet3 = workbook.worksheets[3]
    second_sheet4 = workbook.worksheets[4]
    second_sheet5 = workbook.worksheets[5]
    for i in trange(len(df1)):
        driver = str(df1['司机ID'][i]).split('.')[0]
        name = df1['姓名'][i]
        start_month = df1['出生年月'][i]
        sex = df1['性别'][i]
        ID = df1['身份证号'][i]
        jiguan = df1['籍贯'][i]
        suozaidi = df1['身份证地址'][i]
        chexing = df1['准驾车型'][i]
        phone = df1['司机手机号'][i]
        start_time = df1['驾驶证初领日期'][i]
        # 将name和ID放在元组中
        name_tuple = (name, name, name, name)
        # print(driver, name, ID, riqi)
        # continue
        # 获取单元格的值
        second_sheet['H2'] = driver
        second_sheet['H2'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet['D12'] = name
        second_sheet['D12'].alignment = Alignment(vertical='center', horizontal='center')
        cell_font = Font(underline=None)
        second_sheet['D12'].font = cell_font

        second_sheet['D14'] = ID
        second_sheet['D14'].alignment = Alignment(vertical='center', horizontal='center')
        # second_sheet['H3'] = '陶玉强'
        # second_sheet['H3'].alignment = Alignment(vertical='center', horizontal='center')
        #
        # second_sheet['H4'] = riqi
        # second_sheet['H4'].alignment = Alignment(vertical='center', horizontal='center')

        # second_sheet['D11'] = '海口喜行网络科技有限公司'
        # second_sheet['D11'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B3'] = name
        second_sheet1['B3'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['D3'] = ID
        second_sheet1['D3'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['D4'] = start_month
        second_sheet1['D4'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B5'] = sex
        second_sheet1['B5'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['D5'] = ID
        second_sheet1['D5'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B6'] = jiguan
        second_sheet1['B6'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['D6'] = jiguan
        second_sheet1['D6'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B9'] = chexing
        second_sheet1['B9'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B10'] = suozaidi
        second_sheet1['B10'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['B12'] = phone
        second_sheet1['B12'].alignment = Alignment(vertical='center', horizontal='center')

        second_sheet1['D9'] = start_time
        second_sheet1['D9'].alignment = Alignment(vertical='center', horizontal='center')

        (second_sheet2['B5'].value, second_sheet3['B5'].value, second_sheet4['B5'].value, second_sheet5['B5'].value) = name_tuple
        second_sheet2['H5'] = ID
        second_sheet3['I5'] = ID
        second_sheet4['I5'] = ID
        second_sheet5['F5'] = ID

        workbook.save(rf"C:\Users\XYCX\Desktop\档案汇总\{df1['姓名'][i]}.xlsx")


input_word2()